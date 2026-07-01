import streamlit as st
import pandas as pd
import requests
import re
import io
from docx import Document

# --- 웹사이트 기본 디자인 설정 ---
st.set_page_config(page_title="IB MYP 플래너 및 점검 대시보드", layout="wide", initial_sidebar_state="expanded")

try:
    GOOGLE_SHEET_API_URL = st.secrets["GOOGLE_SHEET_API_URL"]
except Exception:
    GOOGLE_SHEET_API_URL = "선생님의_구글_앱스_스크립트_웹앱_주소_여기에_붙여넣기"

REQUIRED_COLUMNS = ["과목", "세부과목", "유닛명", "주요개념", "관련개념", "세계적맥락", "평가영역"]

st.markdown(
    "<p style='text-align: right; color: #7f8c8d; font-size: 15px; font-weight: bold;'>"
    "🛠️ 제작: 충주미덕중학교 정구성</p>",
    unsafe_allow_html=True,
)
st.title("🏫 IB MYP 유닛 플래너 및 점검 대시보드")
st.markdown("---")

# --- IB MYP 교과 데이터 ---
subject_data = {
    "언어와 문학": {
        "kc": ["의사소통", "창의성", "연결", "관점"],
        "rc": {"일반": ["주요 대상", "등장인물", "맥락", "장르", "상호텍스트성", "시점", "목적", "자기표현", "배경", "구조", "문체", "주제"]}
    },
    "언어 습득": {
        "kc": ["의사소통", "연결", "창의성", "문화"],
        "rc": {"일반": ["발음", "대상 맥락", "관습", "형식", "기능", "의미", "메시지", "유형", "목적", "구조", "단어 선택", "공감", "관용표현", "시점", "주장", "편견", "추론", "문체적 선택", "주제", "어투"]}
    },
    "개인과 사회": {
        "kc": ["변화", "시스템", "세계적 상호작용", "시간·장소 및 공간"],
        "rc": {
            "통합된 인문학": ["인과 관계", "선택", "문화", "형평성", "세계화", "정체성", "혁신과 혁명", "관점", "권력", "과정", "자원", "지속가능성"],
            "역사": ["인과 관계", "문명", "갈등", "협동", "문화", "통치 체제", "정체성", "이념", "혁신과 혁명", "상호 의존성", "관점", "의의"],
            "지리": ["인과 관계", "문화", "격차와 형평성", "다양성", "세계화", "관리와 개입", "연결망", "패턴과 추세", "권력", "과정", "규모", "지속가능성"],
            "경제학": ["선택", "소비", "형평성", "세계화", "성장", "모델", "빈곤", "권력", "자원", "부족", "지속가능성", "무역"],
            "철학": ["타성", "존재와 되기", "신념", "인과 관계", "인간 본성", "정체성", "지식", "자유", "마음/몸", "객관성/주관성", "연결", "가치"],
            "경영관리": ["인과 관계", "경쟁", "협동", "문화", "윤리", "세계화", "혁신", "리더십", "권력", "과정", "전략", "구조"],
            "심리학": ["행동", "유대", "인지", "의식", "발달", "무질서", "집단", "학습", "정신 건강", "정신", "증상", "무의식"],
            "사회학/인류학": ["주체성", "공동체", "문화", "정체성", "제도", "의미", "규범", "사회적 교류", "사회화", "사회적 지위", "구조", "주관성"],
            "정치 과학/시민/정부": ["권위", "시민권", "갈등", "협동", "세계화", "정부", "이념", "통합", "상호 의존성", "리더십", "권력", "권리"],
            "세계종교": ["권위", "신념", "신", "운명", "교리", "도덕", "종교적 감정", "의식과 의례", "신성함", "상징주의", "전통", "예배"]
        }
    },
    "체육과 보건": {
        "kc": ["변화", "의사소통", "관계"],
        "rc": {"일반": ["조정", "균형", "선택", "에너지", "환경", "기능", "상호작용", "운동", "관점", "개선", "공간", "시스템"]}
    },
    "과학": {
        "kc": ["변화", "관계", "시스템"],
        "rc": {
            "통합 과학/생물학": ["균형", "결과", "에너지", "환경", "증거", "형태", "기능", "상호작용", "모델", "운동", "유형", "변형"],
            "화학": ["균형", "조건", "결과", "에너지", "증거", "형태", "기능", "상호작용", "모델", "운동", "유형", "전이"],
            "물리학": ["결과", "발전", "에너지", "환경", "증거", "형태", "기능", "상호작용", "모델", "운동", "유형", "변환"]
        }
    },
    "수학": {
        "kc": ["관계", "형식", "논리"],
        "rc": {"일반": ["변화", "동치성", "일반화", "타당성", "근사", "모델", "패턴", "수량", "표현", "단순화", "공간", "시스템"]}
    },
    "예술": {
        "kc": ["변화", "정체성", "미학"],
        "rc": {
            "공연 예술": ["관객", "경계", "구성", "표현", "장르", "혁신", "해석", "서사", "연극", "형식", "역할", "구조"],
            "시각 예술": ["관객", "경계", "구성", "표현", "장르", "혁신", "해석", "서사", "형식", "묘사", "특징", "시각문화"]
        }
    },
    "디자인": {
        "kc": ["의사소통", "시스템", "공동체", "개발"],
        "rc": {"일반": ["조정", "협력", "인체 공학", "평가", "형태", "기능", "혁신", "발명", "시장과 트렌드", "관점", "자원", "지속가능성"]}
    }
}
global_contexts = ["정체성과 관계성", "시공간적 맥락", "개인적/문화적 표현", "과학과 기술의 혁신", "세계화와 지속가능성", "공정과 발달"]
criteria_list = ["Criterion A", "Criterion B", "Criterion C", "Criterion D"]


def is_google_sheet_connected() -> bool:
    return bool(GOOGLE_SHEET_API_URL) and "여기에" not in GOOGLE_SHEET_API_URL


# ─────────────────────────────────────────
# docx 자동 추출 함수
# ─────────────────────────────────────────
def extract_from_docx(file_bytes: bytes) -> dict:
    """
    IB MYP 유닛 플래너 docx 파일에서 핵심 정보를 추출합니다.
    반환: {"과목", "유닛명", "주요개념", "관련개념", "세계적맥락", "평가영역"} 딕셔너리
    """
    doc = Document(io.BytesIO(file_bytes))
    result = {col: "" for col in REQUIRED_COLUMNS}

    if not doc.tables:
        return result

    # ── 표 0: 과목, 유닛명, 주요개념, 관련개념, 세계적맥락 추출 ──
    table0 = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table0.rows]

    for i, row in enumerate(rows):
        flat = " ".join(row)

        # 과목 & 유닛명: Teacher / Subject group / Unit title 헤더 행 감지
        if "Teacher" in flat and "Subject group" in flat and "Unit title" in flat:
            if i + 1 < len(rows):
                seen = []
                for val in rows[i + 1]:
                    if val and val not in seen:
                        seen.append(val)
                # seen[0]=교사, seen[1]=과목, seen[2]=유닛명
                if len(seen) >= 2:
                    result["과목"] = seen[1]
                if len(seen) >= 3:
                    result["유닛명"] = seen[2]

        # 주요개념 / 관련개념 / 세계적맥락
        if "Key concept" in flat or "핵심 개념" in flat:
            if i + 1 < len(rows):
                data_row = rows[i + 1]
                seen = []
                for val in data_row:
                    if val and val not in seen:
                        seen.append(val)
                if len(seen) >= 1:
                    result["주요개념"] = seen[0]
                if len(seen) >= 2:
                    result["관련개념"] = seen[1]
                if len(seen) >= 3:
                    result["세계적맥락"] = seen[2]

    # ── 표 1: 평가영역(Criterion) 추출 ──
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        for row in table1.rows:
            cell_text = row.cells[0].text if row.cells else ""
            criteria_found = re.findall(r"Criterion\s+[A-D]", cell_text)
            for c in criteria_found:
                c_normalized = c.replace("  ", " ").strip()
                if c_normalized not in result["평가영역"]:
                    if result["평가영역"]:
                        result["평가영역"] += ", "
                    result["평가영역"] += c_normalized

    return result


# ─────────────────────────────────────────
# 데이터 로드 (구글 시트 연동)
# ─────────────────────────────────────────
if "temp_db" not in st.session_state:
    st.session_state.temp_db = pd.DataFrame(columns=REQUIRED_COLUMNS)

if "gs_status" not in st.session_state:
    st.session_state.gs_status = None


@st.cache_data(ttl=30, show_spinner=False)
def fetch_from_google_sheet(url: str):
    response = requests.get(url, timeout=10)
    response.raise_for_status()
    data = response.json()
    df = pd.DataFrame(data)
    for col in REQUIRED_COLUMNS:
        if col not in df.columns:
            df[col] = ""
    return df[REQUIRED_COLUMNS]


def load_data() -> pd.DataFrame:
    if is_google_sheet_connected():
        try:
            df = fetch_from_google_sheet(GOOGLE_SHEET_API_URL)
            st.session_state.gs_status = "connected"
            return df
        except Exception:
            st.session_state.gs_status = "error"
            return st.session_state.temp_db
    st.session_state.gs_status = None
    return st.session_state.temp_db


db_current = load_data()

# 연동 상태 배너
if st.session_state.gs_status == "connected":
    st.success("🟢 구글 시트에 연결되어 실시간으로 저장 중입니다.", icon="✅")
elif st.session_state.gs_status == "error":
    st.warning("🟡 구글 시트 연결 실패. 임시 저장 중입니다.", icon="⚠️")
else:
    st.info("⚪ 구글 시트 미연동. 세션 내 임시 저장 중입니다.", icon="ℹ️")

# ─────────────────────────────────────────
# 사이드바
# ─────────────────────────────────────────
if "menu" not in st.session_state:
    st.session_state.menu = "✨ 유닛 새로 만들기"

st.sidebar.markdown("<h3 style='color: #2c3e50;'>🏫 미덕중 플래너</h3>", unsafe_allow_html=True)
menu_options = ["✨ 유닛 새로 만들기", "📂 파일로 일괄 등록", "📘 개별 과목 오버뷰", "🚨 누락 핵심요소 점검"]

for option in menu_options:
    if st.sidebar.button(option, use_container_width=True):
        st.session_state.menu = option
        st.rerun()

st.sidebar.markdown("---")
st.sidebar.caption(f"현재 등록된 총 유닛 수: {len(db_current)}개")

if is_google_sheet_connected() and st.sidebar.button("🔄 구글 시트 새로고침", use_container_width=True):
    fetch_from_google_sheet.clear()
    st.rerun()

if len(db_current) > 0:
    csv_data = db_current.to_csv(index=False).encode("utf-8-sig")
    st.sidebar.download_button(
        label="📥 전체 데이터 백업 (CSV)",
        data=csv_data,
        file_name="2026_미덕중_IB_MYP_매핑_결과.csv",
        mime="text/csv",
        use_container_width=True,
    )


# ─────────────────────────────────────────
# 유닛 저장 공통 함수
# ─────────────────────────────────────────
def save_unit_to_db(new_data: dict) -> bool:
    saved_to_google = False
    if is_google_sheet_connected():
        try:
            res = requests.post(GOOGLE_SHEET_API_URL, json=new_data, timeout=10)
            if res.status_code == 200:
                saved_to_google = True
                fetch_from_google_sheet.clear()
        except Exception:
            pass

    if not saved_to_google:
        new_row = pd.DataFrame([new_data])
        st.session_state.temp_db = pd.concat(
            [st.session_state.temp_db, new_row], ignore_index=True
        )
        if is_google_sheet_connected():
            st.toast("⚠️ 구글 시트 저장 실패. 임시 저장소에 저장했습니다.", icon="⚠️")
    return saved_to_google


# ==========================================
# 1. 유닛 새로 만들기
# ==========================================
if st.session_state.menu == "✨ 유닛 새로 만들기":
    st.subheader("✨ 새로운 유닛을 설계합니다.")
    selected_subject = st.selectbox(
        "어떤 과목의 유닛을 만드시나요?",
        ["과목을 선택해주세요..."] + list(subject_data.keys())
    )

    if selected_subject != "과목을 선택해주세요...":
        st.markdown("---")
        sub_subject = "일반"
        if len(subject_data[selected_subject]["rc"]) > 1:
            sub_subject = st.selectbox(
                "👉 세부 과목을 선택해주세요",
                list(subject_data[selected_subject]["rc"].keys())
            )

        unit_name = st.text_input("📝 유닛명을 입력해주세요")

        col1, col2 = st.columns(2)
        with col1:
            kc = st.selectbox("🔑 주요 개념 (Key Concept)", subject_data[selected_subject]["kc"])
            gc = st.selectbox("🌍 세계적 맥락 (Global Context)", global_contexts)
        with col2:
            rc = st.multiselect(
                "🔗 관련 개념 (Related Concepts)",
                subject_data[selected_subject]["rc"][sub_subject]
            )
            crit = st.multiselect("🎯 평가 영역 (Criteria)", criteria_list)

        if st.button("💾 저장하고 과목 오버뷰로 이동하기", type="primary", use_container_width=True):
            if not unit_name or len(rc) == 0 or len(crit) == 0:
                st.warning("유닛명, 관련 개념, 평가 영역을 모두 입력해주세요!")
            else:
                new_data = {
                    "과목": selected_subject,
                    "세부과목": sub_subject if sub_subject != "일반" else "없음",
                    "유닛명": unit_name,
                    "주요개념": kc,
                    "관련개념": ", ".join(rc),
                    "세계적맥락": gc,
                    "평가영역": ", ".join(crit),
                }
                save_unit_to_db(new_data)
                st.success("유닛이 안전하게 기록되었습니다!")
                st.session_state.menu = "📘 개별 과목 오버뷰"
                st.rerun()


# ==========================================
# 2. 파일로 일괄 등록 (NEW!)
# ==========================================
elif st.session_state.menu == "📂 파일로 일괄 등록":
    st.subheader("📂 유닛 플래너 파일로 일괄 등록")
    st.markdown(
        "기존에 작성하신 IB MYP 유닛 플래너 **docx 파일**을 업로드하면 "
        "핵심 정보를 자동으로 읽어서 등록해드립니다."
    )
    st.info("💡 여러 파일을 한꺼번에 선택해서 올리면 한 번에 등록됩니다!", icon="ℹ️")

    uploaded_files = st.file_uploader(
        "파일을 여기에 올려주세요 (docx, 여러 개 가능)",
        type=["docx"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        st.markdown("---")
        st.markdown(f"### 📋 총 {len(uploaded_files)}개 파일 미리보기 및 수정")
        st.caption("자동으로 추출한 정보를 확인하고, 잘못된 부분이 있으면 직접 수정 후 등록하세요.")

        extracted_list = []
        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.read()
            extracted = extract_from_docx(file_bytes)
            extracted["_filename"] = uploaded_file.name
            extracted_list.append(extracted)

        # 수정 가능한 미리보기 테이블
        edited_list = []
        all_subjects = list(subject_data.keys())

        for idx, info in enumerate(extracted_list):
            with st.expander(f"📄 {info['_filename']}", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    subj = st.selectbox(
                        "과목", all_subjects,
                        index=all_subjects.index(info["과목"]) if info["과목"] in all_subjects else 0,
                        key=f"subj_{idx}"
                    )
                    unit = st.text_input("유닛명", value=info["유닛명"], key=f"unit_{idx}")
                    kc_val = st.text_input("주요개념", value=info["주요개념"], key=f"kc_{idx}")
                with col2:
                    rc_val = st.text_input("관련개념 (쉼표로 구분)", value=info["관련개념"], key=f"rc_{idx}")
                    gc_val = st.text_input("세계적맥락", value=info["세계적맥락"], key=f"gc_{idx}")
                    crit_val = st.text_input("평가영역 (쉼표로 구분)", value=info["평가영역"], key=f"crit_{idx}")

                edited_list.append({
                    "과목": subj,
                    "세부과목": "없음",
                    "유닛명": unit,
                    "주요개념": kc_val,
                    "관련개념": rc_val,
                    "세계적맥락": gc_val,
                    "평가영역": crit_val,
                })

        st.markdown("---")
        if st.button(
            f"✅ 위 {len(edited_list)}개 유닛 전체 등록하기",
            type="primary",
            use_container_width=True
        ):
            success_count = 0
            fail_list = []
            progress = st.progress(0, text="등록 중...")

            for i, unit_data in enumerate(edited_list):
                if not unit_data["유닛명"]:
                    fail_list.append(f"{i+1}번째 파일 — 유닛명 없음")
                    continue
                save_unit_to_db(unit_data)
                success_count += 1
                progress.progress((i + 1) / len(edited_list), text=f"등록 중... ({i+1}/{len(edited_list)})")

            progress.empty()

            if success_count > 0:
                st.success(f"🎉 총 {success_count}개 유닛이 성공적으로 등록되었습니다!")
            if fail_list:
                st.warning("아래 항목은 등록되지 않았습니다:\n" + "\n".join(fail_list))

            st.session_state.menu = "📘 개별 과목 오버뷰"
            st.rerun()


# ==========================================
# 3. 개별 과목 오버뷰
# ==========================================
elif st.session_state.menu == "📘 개별 과목 오버뷰":
    st.subheader("📘 과목별 유닛 상세 내역")

    subject_to_view = st.selectbox("조회할 과목을 선택하세요", list(subject_data.keys()))
    df_sub = db_current[db_current["과목"] == subject_to_view]

    if len(df_sub) == 0:
        st.info(f"아직 '{subject_to_view}' 과목에 등록된 유닛이 없습니다.")
    else:
        st.dataframe(
            df_sub[["세부과목", "유닛명", "주요개념", "관련개념", "세계적맥락", "평가영역"]],
            use_container_width=True
        )


# ==========================================
# 4. 누락 핵심요소 점검
# ==========================================
elif st.session_state.menu == "🚨 누락 핵심요소 점검":
    st.subheader("🚨 교육과정 누락 요소 점검 대시보드")

    if len(db_current) == 0:
        st.info("아직 등록된 유닛이 없습니다. 유닛을 먼저 생성해주세요.")
    else:
        st.markdown("#### 1️⃣ 전 교과 통합 유닛 목록")
        st.dataframe(
            db_current[["과목", "유닛명", "주요개념", "관련개념", "세계적맥락", "평가영역"]],
            use_container_width=True,
            height=200
        )

        st.markdown("---")
        st.markdown("#### 2️⃣ 과목별 편식(미반영 요소) 점검")

        check_subject = st.selectbox("📌 점검할 과목을 선택하세요", list(subject_data.keys()))

        master_kc = subject_data[check_subject]["kc"]
        master_rc = list(set(
            rc for rc_list in subject_data[check_subject]["rc"].values() for rc in rc_list
        ))

        df_check = db_current[db_current["과목"] == check_subject]
        used_kc = df_check["주요개념"].tolist()
        used_gc = df_check["세계적맥락"].tolist()
        used_rc = [x.strip() for s in df_check["관련개념"] for x in str(s).split(",")]
        used_crit = [x.strip() for s in df_check["평가영역"] for x in str(s).split(",")]

        missing_kc = [k for k in master_kc if k not in used_kc]
        missing_rc = [r for r in master_rc if r not in used_rc]
        missing_gc = [g for g in global_contexts if g not in used_gc]
        missing_crit = [c for c in criteria_list if c not in used_crit]

        st.write("")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("##### 🔑 누락된 **주요 개념**")
            st.error(", ".join(missing_kc)) if missing_kc else st.success("🎉 완벽합니다! 모두 반영되었습니다.")
        with col2:
            st.markdown("##### 🔗 누락된 **관련 개념**")
            st.warning(", ".join(missing_rc)) if missing_rc else st.success("🎉 완벽합니다! 모두 반영되었습니다.")

        st.write("")
        col3, col4 = st.columns(2)
        with col3:
            st.markdown("##### 🌍 누락된 **세계적 맥락**")
            st.error(", ".join(missing_gc)) if missing_gc else st.success("🎉 완벽합니다! 모두 반영되었습니다.")
        with col4:
            st.markdown("##### 🎯 누락된 **평가 영역**")
            st.warning(", ".join(missing_crit)) if missing_crit else st.success("🎉 완벽합니다! 모두 반영되었습니다.")
